"""
Genera el documento de entrega (Word) con:
- Arquitectura definitiva de la app
- Tecnologias definitivas (stack)
- Seccion de capturas de pantalla (para completar manualmente)

Uso:
    python generar_documento_entrega.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

OUTPUT_PATH = "Entrega_Arquitectura_y_Tecnologias.docx"


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def main():
    doc = Document()

    # Estilos base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ------------------------------------------------------------------
    # Portada
    # ------------------------------------------------------------------
    title = doc.add_heading("Guardián de la Vid", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Arquitectura definitiva y tecnologías utilizadas")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "Guardián de la Vid es una aplicación móvil para la detección orientativa "
        "de enfermedades en hojas de vid, desarrollada como parte de una tesis de "
        "grado. Funciona de forma 100% offline mediante un modelo de inteligencia "
        "artificial (TensorFlow Lite) ejecutado directamente en el dispositivo, con "
        "sincronización opcional a la nube cuando hay conectividad."
    )

    doc.add_page_break()

    # ------------------------------------------------------------------
    # 1. Arquitectura definitiva
    # ------------------------------------------------------------------
    doc.add_heading("1. Arquitectura definitiva", level=1)

    doc.add_paragraph(
        "La aplicación sigue una arquitectura por capas, separando claramente la "
        "interfaz de usuario, la lógica de negocio (servicios) y la capa de datos "
        "(modelo de IA local, base de datos SQLite y servicios en la nube)."
    )

    doc.add_heading("1.1 Diagrama de capas", level=2)
    diagram = (
        "┌─────────────────────────────────────────────────────────┐\n"
        "│                     React Native App                     │\n"
        "│  ┌───────────┐  ┌───────────┐  ┌──────────────────────┐ │\n"
        "│  │  Screens  │  │Components │  │      Navigation       │ │\n"
        "│  └─────┬─────┘  └─────┬─────┘  └──────────────────────┘ │\n"
        "│        │              │                                  │\n"
        "│  ┌─────▼──────────────▼──────────────────────────────┐  │\n"
        "│  │                   Services                        │  │\n"
        "│  │  ┌──────────────┐ ┌────────────────┐ ┌──────────┐ │  │\n"
        "│  │  │ tfliteService │ │diagnosisService│ │ imageProc.│ │  │\n"
        "│  │  └──────┬───────┘ └──────┬─────────┘ └────┬─────┘ │  │\n"
        "│  │         │                │                 │       │  │\n"
        "│  │  ┌──────▼──────┐  ┌──────▼──────┐          │       │  │\n"
        "│  │  │ TFLite model │  │  SQLite DB  │          │       │  │\n"
        "│  │  │(assets/model)│  │   (local)   │          │       │  │\n"
        "│  │  └─────────────┘  └─────────────┘          │       │  │\n"
        "│  │                                             │       │  │\n"
        "│  │  jpeg-js → tensor float32 [1, 224, 224, 3] ◄┘       │  │\n"
        "│  └────────────────────────────────────────────────────┘  │\n"
        "│                                                           │\n"
        "│  ┌─────────────────────────────────────────────────────┐ │\n"
        "│  │  Firebase (opcional, solo con conexión a internet)  │ │\n"
        "│  │  Auth  │  Firestore  │  Storage                      │ │\n"
        "│  └─────────────────────────────────────────────────────┘ │\n"
        "└─────────────────────────────────────────────────────────┘"
    )
    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    doc.add_paragraph(
        "La inferencia ocurre dentro del proceso nativo del dispositivo mediante el "
        "módulo react-native-fast-tflite (arquitectura Nitro). No se realiza ninguna "
        "llamada de red durante el diagnóstico."
    )

    doc.add_heading("1.2 Pipeline de diagnóstico", level=2)
    pipeline_steps = [
        ("1. Captura de imagen", "Foto tomada con la cámara o seleccionada de la galería (CameraScreen)."),
        ("2. Preprocesamiento", "imageProcessor.ts redimensiona la imagen a 224×224 px y la comprime al 92% (JPEG)."),
        ("3. Conversión a tensor", "jpegToModelInput.ts decodifica el JPEG, normaliza los píxeles a [0,1], calcula el índice de exceso de verde y evalúa la calidad de imagen (blur, exposición). Resultado: tensor float32 [1,224,224,3]."),
        ("4. Inferencia TFLite", "tfliteService.ts ejecuta el modelo MobileNetV2 (.tflite) sobre el tensor mediante react-native-fast-tflite, obteniendo un vector softmax de 4 probabilidades."),
        ("5. Filtro de escena", "sceneVegetationGate.ts descarta diagnósticos cuando la imagen no corresponde a una hoja de vid, devolviendo 'No es hoja de vid'."),
        ("6. Selección de clase", "predictionFromVector.ts elige la clase de mayor probabilidad y la mapea al nombre de la enfermedad en español."),
        ("7. Nivel de confianza", "confidenceRules.ts clasifica la confianza en 'confiable' (≥85%), 'probable' (65-85%) o 'no concluyente' (<65%)."),
        ("8. Recomendación agronómica", "recommendations.ts asigna nivel de riesgo (Bajo/Moderado/Alto) y una recomendación específica según la enfermedad."),
        ("9. Resultado y persistencia", "ResultScreen.tsx muestra el diagnóstico completo, genera el mapa de calor Grad-CAM y guarda el registro en SQLite (con sincronización opcional a Firebase)."),
    ]
    for titulo, desc in pipeline_steps:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(f"{titulo}: ")
        r1.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "Modo ensemble: el usuario puede capturar hasta 3 imágenes, que se procesan "
        "en paralelo (Promise.all). Los vectores softmax resultantes se promedian "
        "componente a componente antes de continuar con el filtro de escena y la "
        "selección de clase, mejorando la robustez frente a variaciones de ángulo, "
        "iluminación y zoom."
    )

    doc.add_heading("1.3 Modelo de IA", level=2)
    doc.add_paragraph(
        "Arquitectura base: MobileNetV2, preentrenada en ImageNet y ajustada con "
        "transfer learning + fine-tuning sobre un dataset propio de hojas de vid "
        "(1922 imágenes en 4 clases). El modelo se exporta a TensorFlow Lite "
        "(guardian_vid.tflite, ~2.5 MB) y se embebe en assets/model/ junto con "
        "labels.txt y model_metadata.json."
    )
    add_table(
        doc,
        headers=["ID", "Clase", "Enfermedad", "Nivel de riesgo"],
        rows=[
            ["0", "0_sana", "Hoja sana", "Bajo"],
            ["1", "1_mildiu", "Mildiu (Plasmopara viticola)", "Alto"],
            ["2", "2_oidio", "Oídio (Erysiphe necator)", "Moderado"],
            ["3", "3_bacteriana", "Podredumbre bacteriana", "Alto"],
            ["-", "(filtro heurístico)", "No es hoja de vid", "-"],
        ],
    )

    doc.add_heading("1.4 Persistencia y datos", level=2)
    doc.add_paragraph(
        "El historial de diagnósticos se almacena localmente en SQLite (tabla "
        "diagnoses), con migraciones gestionadas mediante un array SQL_MIGRATIONS "
        "que permite evolucionar el esquema sin perder datos. Cada registro incluye "
        "imagen, enfermedad detectada, confianza, nivel de riesgo, recomendación, "
        "fecha, usuario (si hay sesión) y estado de sincronización "
        "(local / pending / synced)."
    )

    doc.add_heading("1.5 Sincronización en la nube (opcional)", level=2)
    doc.add_paragraph(
        "La app es offline-first. Si hay conexión y el usuario tiene sesión activa, "
        "el diagnóstico se sube automáticamente a Firebase (imagen a Storage, "
        "metadatos a Firestore). Si no hay conexión, se guarda localmente con "
        "estado 'pending' y se sincroniza automáticamente cuando NetInfo detecta "
        "que la conectividad se restablece (syncService.ts). Sin sesión iniciada, "
        "el diagnóstico se guarda únicamente en el dispositivo."
    )

    doc.add_page_break()

    # ------------------------------------------------------------------
    # 2. Tecnologías definitivas
    # ------------------------------------------------------------------
    doc.add_heading("2. Tecnologías definitivas", level=1)

    doc.add_heading("2.1 Frontend / Aplicación móvil", level=2)
    add_table(
        doc,
        headers=["Tecnología", "Versión", "Rol"],
        rows=[
            ["React Native", "0.81.5", "Framework de UI multiplataforma (Android / iOS)"],
            ["Expo", "~54.0.33", "Toolchain y gestión de builds nativos"],
            ["TypeScript", "~5.9.2", "Lenguaje principal (modo estricto)"],
            ["React", "19.1.0", "Librería de componentes"],
            ["React Navigation v7", "^7.x", "Navegación entre pantallas (native-stack)"],
        ],
    )

    doc.add_heading("2.2 Machine Learning / Inferencia", level=2)
    add_table(
        doc,
        headers=["Tecnología", "Versión", "Rol"],
        rows=[
            ["react-native-fast-tflite", "^3.0.1", "Módulo nativo (Nitro) para ejecutar modelos .tflite"],
            ["TensorFlow Lite", "—", "Motor de inferencia en el dispositivo"],
            ["jpeg-js", "^0.4.4", "Decodificación JPEG en JS para construir el tensor de entrada"],
            ["expo-image-manipulator", "~14.0.8", "Redimensionado de imágenes a 224×224 px"],
            ["TensorFlow / Keras (entrenamiento)", "TF 2.16 (Python 3.11)", "Entrenamiento y fine-tuning del modelo MobileNetV2"],
        ],
    )

    doc.add_heading("2.3 Persistencia y datos", level=2)
    add_table(
        doc,
        headers=["Tecnología", "Versión", "Rol"],
        rows=[
            ["expo-sqlite", "~16.0.10", "Base de datos local del historial de diagnósticos"],
            ["@react-native-async-storage/async-storage", "^3.1.0", "Almacenamiento clave-valor (preferencias, caché)"],
        ],
    )

    doc.add_heading("2.4 Cámara e imagen", level=2)
    add_table(
        doc,
        headers=["Tecnología", "Versión", "Rol"],
        rows=[
            ["expo-camera", "~17.0.10", "Acceso a la cámara nativa"],
            ["expo-image-picker", "~17.0.10", "Selección de imágenes desde galería"],
        ],
    )

    doc.add_heading("2.5 Conectividad y nube (opcional)", level=2)
    add_table(
        doc,
        headers=["Tecnología", "Versión", "Rol"],
        rows=[
            ["Firebase", "^12.13.0", "Autenticación, Firestore (metadatos) y Storage (imágenes)"],
            ["@react-native-community/netinfo", "^12.0.1", "Detección de conectividad para auto-sincronización"],
        ],
    )

    doc.add_heading("2.6 Herramientas de desarrollo", level=2)
    add_table(
        doc,
        headers=["Tecnología", "Rol"],
        rows=[
            ["Jest + jest-expo", "Tests unitarios"],
            ["Metro bundler", "Bundler de React Native (configurado para assets .tflite)"],
            ["EAS Build", "Sistema de builds nativos de Expo Application Services"],
        ],
    )

    doc.add_page_break()

    # ------------------------------------------------------------------
    # 3. Capturas de pantalla
    # ------------------------------------------------------------------
    doc.add_heading("3. Capturas de pantalla", level=1)
    doc.add_paragraph(
        "Insertar a continuación las capturas de cada pantalla de la aplicación, "
        "siguiendo el flujo de uso: inicio de sesión, pantalla principal, captura "
        "de foto, resultado del diagnóstico (incluyendo mapa de calor Grad-CAM) e "
        "historial de diagnósticos."
    )

    pantallas = [
        "LoginScreen — Inicio de sesión / registro",
        "HomeScreen — Pantalla principal",
        "CameraScreen — Captura de foto (modo simple y modo ensemble)",
        "ResultScreen — Resultado del diagnóstico y mapa de calor Grad-CAM",
        "HistoryScreen — Historial de diagnósticos",
        "DetailScreen — Detalle de un diagnóstico guardado",
    ]
    for pantalla in pantallas:
        doc.add_heading(pantalla, level=2)
        doc.add_paragraph("[ Insertar captura aquí ]")
        doc.add_paragraph()

    doc.save(OUTPUT_PATH)
    print(f"Documento generado: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
